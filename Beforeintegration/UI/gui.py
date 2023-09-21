import tkinter as tk
from tkinter import filedialog
from tkinter import ttk
import subprocess
import csv
from PIL import Image
import torch
from torch import autocast
from diffusers import StableDiffusionPipeline
import docx
from docx import Document
from docx.shared import Inches
from kinki_programfile import kinkizisyo
from kinki_programfile import Bard
import os
import deepl
import glob
from moviepy.editor import *

key=os.environ["DEEPL_AUTH_KEY"]
translator = deepl.Translator(key)
#画像生成プログラム
hugging_token = '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'
ldm = StableDiffusionPipeline.from_pretrained("CompVis/stable-diffusion-v1-4",revision="fp16",torch_dtype=torch.float16,use_auth_token=hugging_token).to("cuda")
def create_background(prompt,numP):
    with autocast("cuda"):
        image = ldm(prompt,
                    negative_prompt="sketches, painting, (character, person, human, people, hands, girl, man, :1.3), user name",
                    height=360,
                    width=640,#解像度をあげるとより具体的な画像が出力される?
                    guidance_scale=10,  # プロンプトの重み（生成画像の類似度（0〜20)）
                    num_inference_steps=50,  # 画像生成に費やすステップ数
                    ).images[0]
    image.save(r"imagefile\image"+ str(numP) +".png")
#prompt = "masterpiece, best quality, ((background only:2)), 構図の指定, 物の指定, 背景の指定, 向きの指定"
#create_background(prompt,numP)→prompt=ポジティブプロンプト,numP何枚目の画像なのか.別画像として保存する際に使用)


# ウィンドウの作成
root = tk.Tk()
root.title("GUI")

# ウィンドウサイズを設定
root.geometry("800x400")  # 幅x高さ

# ウィンドウのサイズ変更を無効にする
root.resizable(width=False, height=False)
root.configure(bg="gray")

button_style_large = {
    "width": 20,
    "height": 10,
    "bg": "lightblue",
    "font": ("Arial", 10),
}

button_style_small = {
    "width": 15,
    "height": 2,
    "bg": "lightblue",
    "font": ("Arial", 10),
}

# メインメニュー画面を表示する関数
def show_main_menu():
    # 現在表示中のフレームを非表示にする
    if current_frame:
        current_frame.grid_remove()

    # メインメニューフレームを作成
    main_menu_frame.grid(row=0, column=0)

# "禁忌チェック"ボタンがクリックされたときの処理
def button1_click():
    global current_frame
    current_frame = check_forbidden_frame

    # メインメニューフレームを非表示にし、禁忌チェック画面を表示
    main_menu_frame.grid_remove()
    check_forbidden_frame.grid(row=0, column=0)

# "ビデオコンテ作成"ボタンがクリックされたときの処理
def button2_click():
    global current_frame
    current_frame = video_create_frame

    # メインメニューフレームを非表示にし、ビデオコンテ作成画面を表示
    main_menu_frame.grid_remove()
    video_create_frame.grid(row=0, column=0)

# "root画面に戻る"ボタンがクリックされたときの処理
def back_to_main_menu():
    global current_frame
    current_frame = main_menu_frame

    # 他のフレームを非表示にし、メインメニューフレームを表示
    check_forbidden_frame.grid_remove()
    video_create_frame.grid_remove()
    data_frame.grid_remove()
    video_data_frame.grid_remove()
    nodata_frame.grid_remove()
    main_menu_frame.grid(row=0, column=0)

def select_file(path):
    file_path = filedialog.askopenfilename(title=path)
    if file_path:
        if path == "ファイルを選択":
            file_path_text.config(state="normal")
            file_path_text.delete(1.0, tk.END)  # テキストボックスをクリア
            file_path_text.insert(tk.END, file_path)
            file_path_text.config(state="disabled")
        elif path == "WORDファイル選択":
            file_path_text1.delete(1.0, tk.END)  # テキストボックスをクリア
            file_path_text1.insert(tk.END, file_path)
            global wordpath
            wordpath = file_path

# 表のデータ
data = []
prompt_filepath = ""
# data.append([1,"なし","なし","なし"])
def execute_action():
    # テキストボックスからファイルパスを取得して禁忌辞書とすり合わせるプログラムを実行
    global data
    global prompt_filepath 
    prompt_filepath = file_path_text.get("1.0",tk.END).replace("\n","")
    global removedtext
    data,removedtext = kinkizisyo.main(prompt_filepath)
    # データをテーブルに追加
    for i in data:
        tree.insert("", "end", values=i)
    global current_frame
    # 表にデータを挿入する
    if data:
        # データを表示するウィンドウを呼び出す
        current_frame = nodata_frame

    elif not data:
        # データを表示するウィンドウを呼び出す
        current_frame = data_frame
    # 他のフレームを非表示にし、表を表示させる画面を表示
    check_forbidden_frame.grid_remove()
    video_create_frame.grid_remove()
    data_frame.grid(row=0, column=0)

def execute_action2():
    # Wordfile内の文字を読み取り、表として表示する
    # result_data = subprocess.run(["python", ""], capture_output=True, text=True)

    # データを表示するウィンドウを呼び出す
    global current_frame
    current_frame = video_data_frame

    # 他のフレームを非表示にし、表を表示させる画面を表示
    check_forbidden_frame.grid_remove()
    video_create_frame.grid_remove()
    video_data_frame.grid(row=0, column=0)

def on_entry_click(event):
    text_widget = event.widget
    if text_widget.get("1.0", tk.END) == "bard APIのPSIDを入力\n":
        text_widget.delete("1.0", tk.END)
        text_widget.config(fg="gray")

def on_focus_out(event):
    text_widget = event.widget 
    if text_widget.get("1.0", tk.END) == "\n":
        text_widget.insert("1.0", "bard APIのPSIDを入力")
        text_widget.config(fg="gray")  


def regenerate_action():
    global data
    global removedtext
    if not bard_api_psid.get("1.0", tk.END)== "bard APIのPSIDを入力\n":
        tmp = bard_api_psid.get("1.0", tk.END).replace("\n","")
        file_path = filedialog.askdirectory(title="ダウンロードファイルを選択")
        if file_path:
            file_path = os.path.join(file_path,"pronpt.txt")
            response = Bard.main(tmp,removedtext)
            with open(file_path,"w") as file:
                file.write(response)
            back_to_main_menu()
   
def Storyboard(wordpath):
    doc = docx.Document(wordpath)
    global tbn
    tbn=len(doc.tables)
    global word_text
    word_text = []
    for i in range(tbn):
        tbl = doc.tables[i]
        tmpW = 0
        for row in tbl.rows:
            row_text = []
            if tmpW == 1:
                for cell in row.cells:
                    row_text.append(cell.text)
                    l_replace = [s.replace("\u3000","") for s in row_text]
                word_text.append(l_replace)
            tmpW = 1
    
    numP = 0
    global icons
    icons = {}
    global secS
    secS=[]
    for i in tree1.get_children():
        tree1.delete(i)
    for i in range(18*tbn):
        if i%3 == 0:
            if word_text[i][5]=="":
                break
            scene = word_text[i][0]
            cut = word_text[i][1]
            content = word_text[i+2][3]
            seconds = word_text[i][5]
            secS.append(seconds)
            kouzu = word_text[i][3].replace("構図：","")
            if kouzu!="":
                kouzu = translator.translate_text(kouzu, target_lang="EN-US").text
            mono = word_text[i][4].replace("物：","")
            if mono!="":
                mono = translator.translate_text(mono, target_lang="EN-US").text
            haikei = word_text[i+1][3].replace("背景：","")
            if haikei!="":
                haikei = translator.translate_text(haikei, target_lang="EN-US").text
            muki = word_text[i+1][4].replace("向き：","")
            if muki!="":
                muki = translator.translate_text(muki, target_lang="EN-US").text
            prompt = "masterpiece, best quality, ((background only:2)), "+kouzu+", "+mono+", "+haikei+", "+muki+""
            #create_background(prompt,numP)
            #im = Image.open(folder_path+"\image"+ str(numP) +".png")
            #back_im = im.copy()
            #back_im = back_im.resize((240, 135))
            #back_im.save(folder_path+r"\resize"+str(numP)+".png", quality=95)
            icons["image"+ str(numP)] = tk.PhotoImage(file=folder_path+r"\resize"+ str(numP) +".png")
            tree1.insert('', 'end',image=icons["image"+ str(numP)],value=(scene, cut,content, seconds),tags=(numP))
            numP = numP + 1
    
    global current_frame
    current_frame = video_data_frame

    # 他のフレームを非表示にし、表を表示させる画面を表示
    check_forbidden_frame.grid_remove()
    video_create_frame.grid_remove()
    video_data_frame.grid(row=0, column=0)

def movie(secS):
    file_list = glob.glob(folder_path+"\image*.png")
    # ファイル名リストを昇順にソート
    file_list.sort()
    # スライドショーを作る元となる静止画情報を格納する処理
    clips = [] 
    tmpMV=0
    for m in file_list:
        #img = Image.open(m)
        clip = ImageClip(m).set_duration(secS[tmpMV])
        clips.append(clip)
        tmpMV= tmpMV+1
    # スライドショーの動画像を作成する処理
    concat_clip = concatenate_videoclips(clips, method="compose")
    concat_clip.write_videofile(r"output.mp4", 
                                fps=24,
                                write_logfile=True,
                                )

def get_selected_item():
    selected_item = tree1.selection()  #選択した行のIDを取得
    if selected_item:
        item = tree1.item(selected_item)
        tags = item['tags'][0] #選択した行のデータを取得
        tags = int(tags)
        if tags!="":
            i=tags*3-2
            kouzu = word_text[i][3].replace("構図：","")
            if kouzu!="":
                kouzu = translator.translate_text(kouzu, target_lang="EN-US").text
            mono = word_text[i][4].replace("物：","")
            if mono!="":
                mono = translator.translate_text(mono, target_lang="EN-US").text
            haikei = word_text[i+1][3].replace("背景：","")
            if haikei!="":
                haikei = translator.translate_text(haikei, target_lang="EN-US").text
            muki = word_text[i+1][4].replace("向き：","")
            if muki!="":
                muki = translator.translate_text(muki, target_lang="EN-US").text
            prompt = "masterpiece, best quality, ((background only:2)), "+kouzu+", "+mono+", "+haikei+", "+muki+""
            create_background(prompt,tags)
            im = Image.open(folder_path+"\image"+ str(tags) +".png")
            back_im = im.copy()
            back_im = back_im.resize((240, 135))
            back_im.save(folder_path+r"\resize"+str(tags)+".png", quality=95)
            
            numP = 0
            global icons
            icons = {}
            for i in tree1.get_children():
                tree1.delete(i)
            for i in range(18*tbn):
                if i%3 == 0:
                    if word_text[i][5]=="":
                        break
                    scene = word_text[i][0]
                    cut = word_text[i][1]
                    content = word_text[i+2][3]
                    seconds = word_text[i][5]
                    icons["image"+ str(numP)] = tk.PhotoImage(file=folder_path+r"\resize"+ str(numP) +".png")
                    tree1.insert('', 'end',image=icons["image"+ str(numP)],value=(scene, cut,content, seconds),tags=(numP))
                    numP = numP + 1
            
def exportDocx():
    movie(secS)
    doc = Document(wordpath)
    file_list = glob.glob(folder_path+r"\resize*.png")
    # ファイル名リストを昇順にソート
    file_list.sort()
    max = len(file_list)
    tmp = 0
    for i in range(tbn):
        table = doc.tables[i]
        for m in range(6):
            if tmp<max:
                row_index = 3*m+1  # 行のインデックス
                column_index = 2  # 列のインデックス
                cell = table.cell(row_index, column_index)
                # 画像を挿入
                cell.paragraphs[0].clear()  # セル内の既存の内容をクリア
                run = cell.paragraphs[0].add_run()
                image_path = file_list[tmp]
                run.add_picture(image_path, width=Inches(2.6), height=Inches(1.4625))
                tmp = tmp + 1
    doc.save(wordpath)
    back_to_main_menu()

# メインメニューフレーム
main_menu_frame = tk.Frame(root, bg="gray")
main_menu_frame.grid(row=0, column=0)

button1 = tk.Button(main_menu_frame, text="禁忌チェック", command=button1_click, **button_style_large)
button1.grid(row=0, column=0, padx=100, pady=100)

button2 = tk.Button(main_menu_frame, text="ビデオコンテ作成", command=button2_click, **button_style_large)
button2.grid(row=0, column=1, padx=100, pady=100)

# 禁忌チェック画面
check_forbidden_frame = tk.Frame(root, bg="gray")

button_back = tk.Button(check_forbidden_frame, text="メイン画面に戻る", command=back_to_main_menu, **button_style_small)
button_back.grid(row=0, column=0, padx=5, pady=5)

button_select_file = tk.Button(check_forbidden_frame, text="テキストファイルを選択", command=lambda:select_file("ファイルを選択"), **button_style_large)
button_select_file.grid(row=1, column=1, padx=200, pady=20)

# テキストボックスを作成
file_path_text = tk.Text(check_forbidden_frame, height=1, width=40, bg="lightgray")
file_path_text.config(state="disabled")
file_path_text.grid(row=2, column=1, padx=10, pady=0)

execute_button1 = tk.Button(check_forbidden_frame, text="実行", command=execute_action, **button_style_small)
execute_button1.grid(row=3, column=1, padx=0, pady=0)


nodata_frame = tk.Frame(root,bg="gray")
nodata_frame_label = tk.Label(nodata_frame, text="検出された放送禁止用語はありませんでした")
nodata_frame_label.grid(row=1, column=1, padx=200, pady=20)
nodata_frame_button = tk.Button(nodata_frame, text="戻る", command=nodata_frame.pack_forget)
nodata_frame_button.grid(row=2, column=1, padx=10, pady=0)

# 表を表示(激うまギャグ)フレーム
data_frame = tk.Frame(root, bg="gray")

# データを表示するためのテーブル（Treeview）を作成
DataFlame_style1 = ttk.Style()
DataFlame_style1.configure('CustomStyle1.Treeview')

columns = ("No.", "禁忌単語", "概要", "改善案")
tree = ttk.Treeview(data_frame, columns=columns, show="headings",style='CustomStyle1.Treeview')

# テーブルの各列にヘッダーを設定
for col in columns:
    tree.heading(col, text=col)
    tree.column(col, width=100)  # 列の幅を調整



# テーブルをフレームに配置
tree.pack(side="top", fill="both", expand=True)
tree.grid(row=1, column=0, padx=200, pady=10)

button_back3 = tk.Button(data_frame, text="メイン画面に戻る", command=back_to_main_menu, **button_style_small)
button_back3.grid(row=0, column=0, padx=5, pady=5,sticky="w")


bard_api_psid = tk.Text(data_frame, height=1, width=40, bg="lightgray",fg="gray")
bard_api_psid.insert("1.0", "bard APIのPSIDを入力")
bard_api_psid.bind("<FocusIn>", on_entry_click)
bard_api_psid.bind("<FocusOut>", on_focus_out)
bard_api_psid.grid(row=2, column=0, padx=5, pady=5)

submit_button = tk.Button(data_frame, text="再生成&ダウンロード",command=regenerate_action,**button_style_small)
submit_button.grid(row=3, column=0, padx=5, pady=5)

# ビデオコンテファイル選択画面
video_create_frame = tk.Frame(root, bg="gray")

button_back2 = tk.Button(video_create_frame, text="メイン画面に戻る", command=back_to_main_menu, **button_style_small)
button_back2.grid(row=0, column=0, padx=5, pady=5)

button_select_video = tk.Button(video_create_frame, text="WORDファイル選択",command=lambda:select_file("WORDファイル選択"), **button_style_large)
button_select_video.grid(row=1, column=1, padx=200, pady=20)

# テキストボックスを作成
file_path_text1 = tk.Text(video_create_frame, height=1, width=40, bg="lightgray")
file_path_text1.grid(row=2, column=1, padx=10, pady=0)

execute_button2 = tk.Button(video_create_frame, text="実行",command=lambda:Storyboard(wordpath),**button_style_small)
execute_button2.grid(row=3, column=1, padx=0, pady=0)

#ビデオ・画像生成用の表を表示する画面
video_data_frame  = tk.Frame(root, bg="gray")

button_back3 = tk.Button(video_data_frame, text="メイン画面に戻る", command=back_to_main_menu, **button_style_small)
button_back3.grid(row=0, column=0, padx=5, pady=5,sticky="w")

button_back3 = tk.Button(video_data_frame, text="選択し再生成",  command=get_selected_item, **button_style_small)
button_back3.grid(row=0, column=0, padx=(0,225), pady=5)

button_back3 = tk.Button(video_data_frame, text="wordファイルに出力",command=exportDocx,**button_style_small)
button_back3.grid(row=0, column=0, padx=0, pady=0,sticky="e")

#treeView
DataFlame_style2 = ttk.Style()
DataFlame_style2.configure('CustomStyle2.Treeview', rowheight=140)
tree1 = ttk.Treeview(video_data_frame, column=('A','B','C','D'), selectmode='browse', height=2,style='CustomStyle2.Treeview')

# 垂直スクロールバー
vsb = ttk.Scrollbar(video_data_frame, orient=tk.VERTICAL, command=tree1.yview)
vsb.grid(row=1, column=1, sticky="news")
tree1.configure(yscrollcommand=vsb.set)
tree1.grid(row=1, column=0, padx=0, pady=0, sticky='nsew')


tree1.heading('#0', text='画面', anchor='center')
tree1.heading('#1', text='シーン', anchor='center')
tree1.heading('#2', text='カット', anchor='center')
tree1.heading('#3', text='内容', anchor='center')
tree1.heading('#4', text='秒', anchor='center')
tree1.column('#0', anchor='center', width=280)
tree1.column('#1', anchor='center', width=40)
tree1.column('#2', anchor='center', width=40)
tree1.column('#3', anchor='center', width=360)
tree1.column('#4', anchor='center', width=50)

current_directory = os.getcwd()
folder_path = current_directory+'\imagefile'
#フォルダが存在しない場合にフォルダを作成
if not os.path.exists(folder_path):
    os.makedirs(folder_path)

tree.grid_rowconfigure(0, weight=1)
tree.grid_columnconfigure(0, weight=1)
# イベントループを開始
current_frame = main_menu_frame
root.mainloop()